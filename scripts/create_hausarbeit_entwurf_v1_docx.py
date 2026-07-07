from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("docs/hausarbeit_entwurf_v1.docx")


def configure(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def p(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.add_run(text)


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.add_run(item)


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), "D0D7DE")
        tbl_borders.append(tag)
    tbl_pr.append(tbl_borders)


def cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.font.size = Pt(9.5)
    run.bold = bold


def add_table(doc: Document, headers: list[str], rows: list[tuple[str, ...]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    borders(table)
    for cell, header in zip(table.rows[0].cells, headers):
        shade(cell, "F2F4F7")
        cell_text(cell, header, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            cell_text(cell, value)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.add_paragraph()


def add_title(doc: Document) -> None:
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run("Hybrid AI Agent for Dynamic Used Car Pricing")
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("0B2545")

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(10)
    run = subtitle.add_run("Erster vollständiger Entwurf der Hausarbeit | Team MAIL | Stand: 28.06.2026")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string("555555")


def add_note(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    borders(table)
    cell = table.cell(0, 0)
    shade(cell, "FFF7E6")
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(title)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string("7A5A00")
    paragraph = cell.add_paragraph(body)
    paragraph.paragraph_format.space_after = Pt(0)
    doc.add_paragraph()


def main() -> None:
    doc = Document()
    configure(doc)
    add_title(doc)

    doc.add_heading("Abstract", level=1)
    p(
        doc,
        "Diese Arbeit entwickelt und evaluiert einen hybriden AI Agent zur dynamischen "
        "Gebrauchtwagenpreisgestaltung. Ausgangspunkt ist die Beobachtung, dass Gebrauchtwagenpreise "
        "nicht nur vom einzelnen Fahrzeug abhängen, sondern zugleich durch makroökonomische "
        "Marktbewegungen und saisonale Nachfrageeffekte beeinflusst werden. Das System kombiniert "
        "deshalb drei bewusst getrennte Komponenten: Stage 1 schätzt mit einem Machine-Learning-Modell "
        "den fahrzeugbezogenen Basiswert, Stage 2 skaliert diesen Wert über einen CPI-Multiplikator auf "
        "das gewünschte Marktpreisniveau, und Stage 3 ergänzt eine konservative saisonale Anpassung nach "
        "Karosserieform und Monat. Auf einem gemeinsamen Testsplit reduziert das Stage-1-V2-Modell den "
        "mittleren absoluten Fehler gegenüber dem ursprünglichen Modell von 1.830,95 USD auf 1.370,15 USD, "
        "was einer Verbesserung von 25,17 % entspricht. Die CPI-Anpassung zeigt im basisnahen historischen "
        "Testzeitraum 2014–2015 nur geringe Veränderungen, ist aber für Projektionen in spätere Marktphasen "
        "relevant. Die saisonale Regel verbessert den MAE auf einem getrennten Regel-Holdout von 1.353,15 USD "
        "auf 1.339,84 USD. Die Ergebnisse zeigen, dass eine modulare Architektur sowohl Genauigkeit als auch "
        "Interpretierbarkeit unterstützt."
    )

    doc.add_heading("1. Introduction", level=1)
    p(
        doc,
        "Der Markt für Gebrauchtwagen ist durch eine Kombination aus hoher wirtschaftlicher Relevanz, "
        "starker Produktheterogenität und ausgeprägter Informationsasymmetrie geprägt. Bereits Akerlof "
        "(1970) nutzte den Gebrauchtwagenmarkt als zentrales Beispiel, um zu zeigen, wie Qualitätsunsicherheit "
        "zwischen Verkäufern und Käufern zu ineffizienten Marktmechanismen führen kann. Für Käufer ist der "
        "tatsächliche Zustand eines Fahrzeugs vor dem Kauf häufig nur unvollständig beobachtbar; für Verkäufer "
        "besteht gleichzeitig ein Anreiz, Fahrzeuge möglichst wertstabil darzustellen. Datenbasierte "
        "Bewertungsmodelle können diese Informationslücke nicht vollständig schließen, sie können aber "
        "beobachtbare Merkmale systematisch verarbeiten und dadurch eine transparentere Preisindikation liefern."
    )
    p(
        doc,
        "Klassische Bewertungslogiken folgen häufig einer hedonischen Perspektive. Danach ergibt sich der Preis "
        "eines Gutes nicht nur aus dem Produkt als Ganzem, sondern aus einem Bündel einzelner Eigenschaften. "
        "Rosen (1974) beschreibt Produktmärkte als Märkte impliziter Preise, auf denen Eigenschaften wie Qualität, "
        "Ausstattung oder technische Merkmale in den beobachtbaren Marktpreis einfließen. Für Fahrzeuge bedeutet "
        "dies, dass Marke, Modell, Laufleistung, Alter, Zustand, Karosserieform und Ausstattung gemeinsam betrachtet "
        "werden müssen. Moderne Machine-Learning-Verfahren erweitern diese Logik, weil sie nichtlineare Zusammenhänge "
        "und Interaktionen zwischen Merkmalen erfassen können."
    )
    p(
        doc,
        "Gleichzeitig reichen fahrzeugbezogene Merkmale allein nicht aus, um dynamische Gebrauchtwagenpreise zu "
        "erklären. Der US-Gebrauchtwagenmarkt wurde insbesondere seit 2020 durch makroökonomische Verwerfungen, "
        "Lieferkettenprobleme und veränderte Angebots- und Nachfragestrukturen beeinflusst. Ein statisches Modell, "
        "das ausschließlich auf historischen Auktionsdaten aus 2014 und 2015 trainiert wurde, würde solche späteren "
        "Preisniveauverschiebungen nicht automatisch abbilden. Deshalb wird in dieser Arbeit ein externer "
        "Preisindex genutzt, um den fahrzeugbezogenen Basiswert auf ein aktuelles oder frei wählbares Marktpreisniveau "
        "zu skalieren."
    )
    p(
        doc,
        "Die vorliegende Arbeit entwickelt einen hybriden AI Agent zur dynamischen Gebrauchtwagenpreisgestaltung. "
        "Die Architektur ist modular aufgebaut: Stage 1 bestimmt den zeitneutralen Fahrzeugwert, Stage 2 passt diesen "
        "Wert über einen CPI-Multiplikator an das Marktpreisniveau an, und Stage 3 ergänzt eine saisonale Feinkorrektur. "
        "Die zentrale Forschungsfrage lautet, ob eine solche modulare Kombination aus Machine Learning, makroökonomischer "
        "Indexierung und regelbasierter Saisonalität eine präzisere und zugleich interpretierbare Preisindikation erzeugen kann."
    )

    doc.add_heading("2. Theoretical Background", level=1)
    doc.add_heading("2.1 Information Asymmetry in Used-Car Markets", level=2)
    p(
        doc,
        "Der Gebrauchtwagenmarkt ist ein klassischer Fall asymmetrischer Information. Verkäufer kennen die Historie, "
        "Pflege und mögliche verborgene Mängel eines Fahrzeugs häufig besser als Käufer. Akerlof (1970) zeigt, dass "
        "solche Informationsvorteile zu adverser Selektion führen können: Wenn Käufer die Qualität eines Fahrzeugs "
        "nicht zuverlässig erkennen, orientieren sie sich an einer durchschnittlichen Qualitätserwartung. Dadurch kann "
        "der Preis für hochwertige Fahrzeuge zu niedrig ausfallen, während minderwertige Fahrzeuge relativ attraktiv "
        "zu verkaufen sind. Für ein Pricing-System folgt daraus, dass beobachtbare Qualitätsindikatoren wie Zustand, "
        "Laufleistung und Alter explizit berücksichtigt werden sollten."
    )
    doc.add_heading("2.2 Hedonic Pricing", level=2)
    p(
        doc,
        "Die hedonische Preistheorie liefert die ökonomische Grundlage für die Modellierung eines Fahrzeugpreises als "
        "Funktion mehrerer Attribute. Rosen (1974) argumentiert, dass beobachtete Marktpreise implizite Preise einzelner "
        "Produkteigenschaften enthalten. Für die vorliegende Arbeit bedeutet dies, dass Stage 1 den Fahrzeugwert nicht "
        "als pauschalen Durchschnittspreis schätzt, sondern als Ergebnis eines Merkmalsbündels. Dazu zählen numerische "
        "Variablen wie Fahrzeugalter, Laufleistung und Zustand sowie kategoriale Variablen wie Marke, Modell, Ausstattung, "
        "Karosserieform, Getriebe, Bundesstaat, Außenfarbe und Innenfarbe."
    )
    doc.add_heading("2.3 Machine Learning for Tabular Price Prediction", level=2)
    p(
        doc,
        "Used-Car-Pricing ist ein typisches Regressionsproblem auf tabellarischen Daten. Pal et al. (2017) zeigen anhand "
        "eines Random-Forest-Ansatzes, dass überwachte Lernverfahren zur Schätzung von Gebrauchtwagenpreisen eingesetzt "
        "werden können. Neuere Ansätze berücksichtigen zusätzlich Unsicherheit in den Prognosen; Madhusudhanan et al. "
        "(2024) schlagen beispielsweise ein probabilistisches Modell für Used-Car-Pricing vor. In dieser Arbeit wird "
        "für Stage 1 ein XGBoost-basiertes Ensemble genutzt. XGBoost ist ein skalierbares Tree-Boosting-System, das für "
        "große tabellarische Datensätze geeignet ist und in vielen Machine-Learning-Wettbewerben starke Ergebnisse erzielt "
        "(Chen & Guestrin, 2016)."
    )
    doc.add_heading("2.4 Price Indices and Market-Level Adjustment", level=2)
    p(
        doc,
        "Stage 2 basiert auf der Idee, dass ein fahrzeugbezogener Basiswert an das allgemeine Marktpreisniveau angepasst "
        "werden muss. Die U.S. Bureau of Labor Statistics beschreibt den Used Cars and Trucks Index als Bestandteil des "
        "Consumer Price Index, der gebrauchte Fahrzeuge im Alter von zwei bis sieben Jahren umfasst (U.S. Bureau of Labor "
        "Statistics, n.d.). Über FRED wird diese Reihe monatlich als offizieller Index bereitgestellt (U.S. Bureau of Labor "
        "Statistics, 2026). Methodisch wird der Index auf das Jahr 2015 normiert und als Multiplikator verwendet. Dadurch "
        "bleibt Stage 1 für den relativen Fahrzeugwert zuständig, während Stage 2 den zeitlichen Marktpreislevel-Effekt "
        "abbildet."
    )
    doc.add_heading("2.5 Seasonality and Rule-Based Hybrid Systems", level=2)
    p(
        doc,
        "Saisonale Nachfrageeffekte können besonders bei Fahrzeugen plausibel sein, weil unterschiedliche Karosserieformen "
        "zu unterschiedlichen Jahreszeiten attraktiver erscheinen. Cabriolets können beispielsweise in wärmeren Monaten "
        "stärker nachgefragt werden, während SUVs oder allradnahe Fahrzeuge vor Winterperioden relevanter werden können. "
        "Gleichzeitig ist methodisch Vorsicht geboten: Rohe Monatsdurchschnitte können verzerrt sein, wenn in einzelnen "
        "Monaten andere Fahrzeugtypen, Altersstrukturen oder Zustände verkauft wurden. Die hier entwickelte Stage-3-Regel "
        "berechnet Saisonfaktoren deshalb aus CPI-normalisierten und fahrzeugmixbereinigten Modellabweichungen. Stage 3 "
        "ist damit keine zweite künstliche Intelligenz, sondern eine konservative regelbasierte Korrekturschicht."
    )

    doc.add_heading("3. Data and Preprocessing", level=1)
    p(
        doc,
        "Die Micro-Datenbasis des Projekts besteht aus US-amerikanischen Auktionsdaten des Manheim Used Car Auction "
        "Datensatzes. Der bereinigte Datensatz umfasst 558.743 Verkäufe aus den Jahren 2014 und 2015. Für das strenge "
        "Stage-1-V1/V2-Neutraining wurden 529.169 bereinigte Zeilen verwendet, von denen 423.335 in das Training und "
        "105.834 in den Testsplit eingingen. Die Daten enthalten unter anderem Baujahr, Marke, Modell, Ausstattungsvariante, "
        "Karosserieform, Getriebe, Bundesstaat, Zustand, Kilometer- bzw. Meilenstand, Außenfarbe, Innenfarbe und den realen "
        "Verkaufspreis."
    )
    p(
        doc,
        "Ein wichtiger methodischer Schritt ist die Vermeidung von Target Leakage. Deshalb wurden VIN, Verkäufer und der "
        "Manheim Market Report Wert nicht als Modellfeatures verwendet. Besonders MMR wäre problematisch, weil es selbst "
        "bereits eine externe Preisbewertung darstellt und die Modellleistung künstlich erhöhen könnte. Stattdessen nutzt "
        "Stage 1 ausschließlich Merkmale, die aus Nutzersicht plausibel als Fahrzeuginformationen angegeben werden können."
    )
    p(
        doc,
        "Für die makroökonomische Ebene wurde ein eigener Indexdatensatz aufgebaut. Dieser enthält monatliche Werte für "
        "den Gebrauchtwagen-CPI sowie weitere makroökonomische Kontextvariablen wie Federal Funds Rate, Konsumentenstimmung, "
        "Arbeitslosenquote und High-Yield-Spread. Für das aktive Preissystem wird jedoch bewusst nur der CPI-Multiplikator "
        "direkt in Stage 2 verwendet. Weitere Makrosignale dienen vorerst der Interpretation und einer späteren LLM-basierten "
        "Erklärungsschicht."
    )

    doc.add_heading("4. Methodology: Three-Stage Architecture", level=1)
    p(
        doc,
        "Die Kernidee der Methodik ist die strikte Trennung von Fahrzeugwert, Marktpreisniveau und Saisonalität. Dadurch "
        "wird verhindert, dass das System verschiedene Effekte in einem einzigen schwer interpretierbaren Modell vermischt. "
        "Der finale Preis ergibt sich formal als Produkt aus Stage-1-Basiswert, Stage-2-CPI-Multiplikator und Stage-3-"
        "Saisonfaktor:"
    )
    p(doc, "Final Price = Stage-1-Basiswert × CPI-Multiplikator × Saisonfaktor.")
    add_table(
        doc,
        ["Stage", "Aufgabe", "Technische Umsetzung", "Output"],
        [
            ("Stage 1", "Fahrzeugwert", "XGBoost-Ensemble mit Fahrzeugfeatures", "Zeitneutraler Basispreis"),
            ("Stage 2", "Marktpreisniveau", "CPI-Multiplikator relativ zu 2015", "Marktangepasster Preis"),
            ("Stage 3", "Saisonalität", "Regelbasierte Faktoren nach Karosserie und Monat", "Finaler Preis und Verkaufshinweis"),
        ],
    )

    doc.add_heading("4.1 Stage 1: Vehicle Value Model", level=2)
    p(
        doc,
        "Stage 1 ist das eigentliche Machine-Learning-Modell. Die ursprüngliche Version V1 nutzte einen "
        "HistGradientBoostingRegressor mit einer begrenzteren Feature-Auswahl. Im Projektverlauf wurde zusätzlich ein "
        "Stage-1-V2-Modell entwickelt. Dieses V2-Modell ist ein 50/50-Ensemble aus zwei XGBoost-Modellen: Eine Komponente "
        "prognostiziert den Preis direkt in Dollar, die andere den logarithmierten Preis. Die Kombination soll sowohl "
        "absolute Preisfehler reduzieren als auch unterschiedliche Preisbereiche stabiler abdecken."
    )
    p(
        doc,
        "V2 nutzt zusätzliche Fahrzeugmerkmale wie Ausstattungsvariante, Getriebe, Bundesstaat, Außenfarbe, Innenfarbe "
        "und eine explizite Marke-Modell-Interaktion. Gleichzeitig enthält V2 bewusst keinen Verkaufsmonat und kein "
        "year_month-Feature. Diese Entscheidung ist zentral für die Architektur: Zeitliche Marktbewegungen gehören zu "
        "Stage 2, saisonale Monatslogik zu Stage 3. Der starke Genauigkeitsgewinn von V2 ist deshalb ein Stage-1-Ergebnis "
        "und darf nicht als Wirkung der Saisonalität interpretiert werden."
    )

    doc.add_heading("4.2 Stage 2: CPI Macro Adjustment", level=2)
    p(
        doc,
        "Stage 2 multipliziert den von Stage 1 geschätzten Basiswert mit einem CPI-Multiplikator. Dieser Multiplikator ist "
        "auf den Jahresdurchschnitt 2015 normiert. Ein Wert von 1,2177 bedeutet beispielsweise, dass das allgemeine "
        "Preisniveau des zugrunde liegenden Gebrauchtwagenindex rund 21,77 % über dem Basisniveau von 2015 liegt. Stage 2 "
        "modelliert damit keine individuelle Fahrzeugqualität, sondern die makroökonomische Preisverschiebung des Marktes."
    )

    doc.add_heading("4.3 Stage 3: Seasonal Adjustment", level=2)
    p(
        doc,
        "Stage 3 ergänzt den Stage-2-Preis um eine saisonale Feinkorrektur. Die Faktoren werden nach Karosserieform und "
        "Verkaufsmonat berechnet. Dabei werden nicht rohe Monatsdurchschnitte genutzt, weil solche Durchschnittswerte durch "
        "unterschiedliche Fahrzeugmixe verzerrt sein könnten. Stattdessen vergleicht die Methode CPI-normalisierte Verkaufspreise "
        "mit den Vorhersagen des zeitneutralen Stage-1-V2-Modells. Die verbleibenden systematischen Abweichungen werden als "
        "vorsichtiger saisonaler Faktor interpretiert."
    )
    p(
        doc,
        "Um Überanpassung zu vermeiden, werden alle Saisonfaktoren Richtung 1,0 geglättet und auf den Bereich 0,85 bis 1,15 "
        "begrenzt. Monate ohne historische Daten bleiben neutral bei 1,0. Eine Best- oder Worst-Month-Empfehlung wird nur "
        "angezeigt, wenn mindestens zwei Monate mit jeweils wenigstens 100 Beobachtungen für die betreffende Karosserieform "
        "vorhanden sind."
    )

    doc.add_heading("5. Implementation", level=1)
    p(
        doc,
        "Die Implementierung ist als reproduzierbare Python-Pipeline aufgebaut. Die wichtigsten Skripte liegen im Ordner "
        "scripts/. Stage 1 wird über train_stage1_v2.py trainiert und über stage1_runtime.py in der App geladen. Stage 2 "
        "wird durch stage2_macro.py umgesetzt, das den CPI-Multiplikator aus macro_index.csv liest und auf den Basispreis "
        "anwendet. Stage 3 wird durch stage3_seasonality.py umgesetzt und über evaluate_stage3.py evaluiert. Die Streamlit-"
        "App kombiniert alle drei Stufen zu einer interaktiven Demo."
    )
    p(
        doc,
        "Die App ist bewusst nutzerorientiert gestaltet. Sie verwendet deutsche Labels, nimmt Kilometer entgegen und rechnet "
        "sie intern in Meilen um, da der Trainingsdatensatz US-basiert ist. Farben, Bundesstaaten, Getriebe und Karosserieformen "
        "werden im Interface verständlich dargestellt, während im Hintergrund weiterhin die englischen Modellkategorien genutzt "
        "werden. Zusätzlich wurde die Karosserieform an Marke und Modell gekoppelt, damit unrealistische Kombinationen wie ein "
        "nicht existierender Kombi-Supersportwagen nicht angeboten werden."
    )
    p(
        doc,
        "Für normale Nutzer zeigt die App nur die wichtigsten Preis- und Verkaufsinformationen. Technische Details wie "
        "Preisaufbau, Modellgüte, Backtests und makroökonomische Kontextdaten sind hinter Toggles versteckt. Diese Trennung "
        "unterstützt die Verständlichkeit der Demo und erlaubt dem Projektteam gleichzeitig, die wissenschaftlichen Details "
        "bei Bedarf sichtbar zu machen."
    )

    doc.add_heading("6. Evaluation and Results", level=1)
    doc.add_heading("6.1 Stage 1 Results", level=2)
    p(
        doc,
        "Für den methodisch bevorzugten V1/V2-Vergleich wurden beide Modelle von Grund auf auf exakt demselben Split neu "
        "trainiert. Dadurch wird verhindert, dass Unterschiede durch abweichende Trainings- oder Testdaten entstehen. Der "
        "Testsplit umfasst 105.834 zuvor unangetastete Zeilen. V2 reduziert den MAE von 1.830,95 USD auf 1.370,15 USD. "
        "Dies entspricht einer absoluten Verbesserung von 460,80 USD und einer relativen MAE-Reduktion von 25,17 %. Das "
        "95%-Bootstrap-Intervall der Verbesserung liegt zwischen 450,74 USD und 470,83 USD."
    )
    add_table(
        doc,
        ["Modell", "MAE", "RMSE", "R²", "MAPE"],
        [
            ("V1 HistGradientBoosting", "1.830,95 USD", "3.276,81 USD", "0,8818", "16,45 %"),
            ("V2 XGBoost-Ensemble", "1.370,15 USD", "2.400,34 USD", "0,9366", "15,13 %"),
        ],
    )

    doc.add_heading("6.2 Stage 2 Results", level=2)
    p(
        doc,
        "Im historischen Backtest verändert Stage 2 die Genauigkeit nur geringfügig. Der MAE steigt von 1.370,16 USD auf "
        "1.376,22 USD. Dieser Effekt ist erwartbar, weil die Testdaten aus 2014–2015 nahe am CPI-Basisjahr 2015 liegen "
        "und die Multiplikatoren im Testset entsprechend nahe bei 1,0 liegen. Der praktische Nutzen von Stage 2 liegt "
        "daher weniger in einer Verbesserung des historischen MAE, sondern in der Fähigkeit, einen Basispreis auf spätere "
        "Marktphasen zu übertragen."
    )
    add_table(
        doc,
        ["Metrik", "Referenz-Baseline", "Mit Stage 2", "Änderung"],
        [
            ("MAE", "1.370,16 USD", "1.376,22 USD", "+6,06 USD"),
            ("RMSE", "2.400,34 USD", "2.411,11 USD", "leicht höher"),
            ("R²", "0,9366", "0,9360", "-0,0006"),
            ("MAPE", "15,13 %", "15,11 %", "nahezu gleich"),
        ],
    )
    p(
        doc,
        "Die Vorwärtsprojektion zeigt dagegen deutliche Preisniveauverschiebungen. Für 2026-06 ergibt der gespeicherte "
        "Makrostand einen Multiplikator von 1,2177. Ein Stage-1-Basiswert von 12.495 USD würde dadurch auf etwa 15.216 USD "
        "steigen. Diese Größenordnung illustriert den COVID-bedingten und danach teilweise persistierenden Preisanstieg "
        "im Gebrauchtwagenmarkt."
    )

    doc.add_heading("6.3 Stage 3 Results", level=2)
    p(
        doc,
        "Stage 3 wurde getrennt als saisonale Regel evaluiert. Die Faktoren wurden aus 80 % der regelrelevanten Daten "
        "berechnet und auf den übrigen 105.688 Zeilen geprüft. Dabei sinkt der MAE auf CPI-normalisierten Preisen von "
        "1.353,15 USD auf 1.339,84 USD, was einer Verbesserung von 0,98 % entspricht. Der Effekt ist bewusst kleiner als "
        "die Stage-1-V2-Verbesserung, aber methodisch plausibel: Saisonalität soll nur eine Feinkorrektur sein."
    )
    add_table(
        doc,
        ["Karosserie", "Beobachtungen", "Bester Monat", "Bester Effekt", "Schwächster Monat", "Schwächster Effekt"],
        [
            ("sedan", "233.725", "März", "+3,2 %", "Juni", "-2,6 %"),
            ("suv", "139.801", "März", "+2,0 %", "April", "-4,5 %"),
            ("hatchback", "25.620", "Januar", "+2,1 %", "Juni", "-4,5 %"),
            ("convertible", "10.102", "März", "+1,3 %", "Dezember", "-1,6 %"),
        ],
    )
    p(
        doc,
        "Wichtig ist die Datenabdeckung: Die historischen Daten enthalten Verkäufe für Januar bis Juli und Dezember, "
        "aber keine Verkäufe für August bis November. Für diese fehlenden Monate bleibt der Saisonfaktor neutral bei 1,0. "
        "Diese konservative Entscheidung verhindert, dass das System Empfehlungen auf Basis nicht vorhandener Daten erzeugt."
    )

    doc.add_heading("7. Discussion", level=1)
    p(
        doc,
        "Die Ergebnisse zeigen, dass die größte quantitative Verbesserung aus Stage 1 stammt. Das V2-Modell nutzt mehr "
        "Fahrzeugmerkmale und eine stärkere Modellarchitektur, wodurch der MAE deutlich sinkt. Dieser Befund passt zur "
        "hedonischen Grundidee, dass Fahrzeugpreise aus vielen Eigenschaften zusammengesetzt sind. Die Erweiterung um "
        "Ausstattung, Getriebe, Standort und Farben erlaubt eine genauere Unterscheidung ähnlicher Fahrzeuge."
    )
    p(
        doc,
        "Stage 2 verbessert den historischen Backtest nicht, erfüllt aber eine andere Aufgabe. Da die Trainings- und Testdaten "
        "nahe am Basisjahr liegen, kann die CPI-Korrektur im Rückblick nur geringe Veränderungen bewirken. Für eine Anwendung, "
        "die Preise in späteren Jahren oder im aktuellen Marktumfeld berechnen soll, ist diese Stufe dennoch wichtig. Ohne "
        "Stage 2 würde ein Modell, das auf 2014–2015-Daten trainiert wurde, strukturell zu niedrige Preise für Perioden mit "
        "höherem Gebrauchtwagenpreisniveau liefern."
    )
    p(
        doc,
        "Stage 3 liefert nur eine kleine Verbesserung, was aber positiv zu interpretieren ist. Würde die saisonale Regel sehr "
        "große Effekte zeigen, bestünde die Gefahr, dass Fahrzeugmix oder Datenlücken fälschlich als Saisonalität interpretiert "
        "werden. Die gewählte Methode kontrolliert bereits für CPI und Stage-1-Fahrzeugmerkmale und glättet unsichere Faktoren "
        "Richtung 1,0. Damit bleibt Stage 3 konservativ und nachvollziehbar."
    )

    doc.add_heading("8. Limitations", level=1)
    bullets(
        doc,
        [
            "Die Daten stammen aus US-Auktionen und sind nicht ohne Weiteres auf europäische Retailmärkte übertragbar.",
            "Der Datensatz deckt überwiegend 2014–2015 ab und enthält keine vollständige Jahresabdeckung für alle Monate.",
            "Der V1/V2-Vergleich verwendet einen Random Split, aber keinen echten temporalen Zukunftstest.",
            "Stage 2 nutzt einen allgemeinen CPI-Multiplikator und modelliert keine segmentspezifischen Marktverschiebungen.",
            "Stage 3 ist eine regelbasierte Heuristik und kein kausaler Nachweis saisonaler Zahlungsbereitschaft.",
            "Die App ist ein Prototyp und ersetzt keine professionelle Fahrzeugbegutachtung.",
        ],
    )

    doc.add_heading("9. Conclusion and Outlook", level=1)
    p(
        doc,
        "Die Arbeit zeigt, dass ein hybrider AI Agent für dynamische Gebrauchtwagenpreisgestaltung technisch und methodisch "
        "sinnvoll umgesetzt werden kann. Die modulare Architektur trennt Fahrzeugwert, Marktpreisniveau und Saisonalität "
        "klar voneinander. Dadurch wird das System besser interpretierbar als ein einzelnes End-to-End-Modell, ohne auf die "
        "Prognosekraft moderner Machine-Learning-Verfahren zu verzichten."
    )
    p(
        doc,
        "Der wichtigste quantitative Fortschritt liegt in Stage 1 V2. Das XGBoost-Ensemble reduziert den MAE gegenüber dem "
        "ursprünglichen Modell um 25,17 %. Stage 2 ergänzt diese Schätzung um eine ökonomisch begründete Marktpreisanpassung, "
        "während Stage 3 eine kleine, vorsichtige Saisonkorrektur hinzufügt. Zusammen entsteht ein System, das nicht nur einen "
        "Preis ausgibt, sondern den Weg zu diesem Preis nachvollziehbar macht."
    )
    p(
        doc,
        "Für zukünftige Arbeiten bieten sich mehrere Erweiterungen an. Erstens sollte ein echter temporaler Holdout genutzt "
        "werden, um die Prognosefähigkeit für spätere Jahre strenger zu testen. Zweitens könnten segmentspezifische CPI- oder "
        "Marktindikatoren eingeführt werden. Drittens wäre eine LLM-Orchestrierung sinnvoll, die die numerischen Ergebnisse "
        "in eine natürlichsprachliche, nutzerangepasste Erklärung überführt. Schließlich könnte das System durch weitere "
        "Datenquellen, etwa regionale Retailpreise oder aktuelle Angebotsdaten, verbessert werden."
    )

    add_note(
        doc,
        "Prüfhinweis vor finaler Abgabe",
        "Für den offiziellen Used-Cars-and-Trucks-CPI führt FRED die Serie CUSR0000SETA02. "
        "Falls Code oder ältere Dokumentation CUSR0000SETA01 nennen, sollte das vor der finalen Hausarbeit abgeglichen "
        "und gegebenenfalls korrigiert werden.",
    )

    doc.add_heading("References", level=1)
    references = [
        "Akerlof, G. A. (1970). The market for “lemons”: Quality uncertainty and the market mechanism. The Quarterly Journal of Economics, 84(3), 488–500. https://doi.org/10.2307/1879431",
        "Chen, T., & Guestrin, C. (2016). XGBoost: A scalable tree boosting system. In Proceedings of the 22nd ACM SIGKDD International Conference on Knowledge Discovery and Data Mining (pp. 785–794). Association for Computing Machinery. https://doi.org/10.1145/2939672.2939785",
        "Madhusudhanan, K., Behrens, G., Stubbemann, M., & Schmidt-Thieme, L. (2024). ProbSAINT: Probabilistic tabular regression for used car pricing. arXiv. https://arxiv.org/abs/2403.03812",
        "Pal, N., Arora, P., Sundararaman, D., Kohli, P., & Palakurthy, S. S. (2017). How much is my car worth? A methodology for predicting used cars prices using Random Forest. arXiv. https://arxiv.org/abs/1711.06970",
        "Rosen, S. (1974). Hedonic prices and implicit markets: Product differentiation in pure competition. Journal of Political Economy, 82(1), 34–55. https://doi.org/10.1086/260169",
        "U.S. Bureau of Labor Statistics. (n.d.). Measuring price change in the CPI: Used cars and trucks. U.S. Department of Labor. https://www.bls.gov/cpi/factsheets/used-cars-and-trucks.htm",
        "U.S. Bureau of Labor Statistics. (2026). Consumer Price Index for All Urban Consumers: Used Cars and Trucks in U.S. City Average [CUSR0000SETA02]. FRED, Federal Reserve Bank of St. Louis. https://fred.stlouisfed.org/series/CUSR0000SETA02",
    ]
    for ref in references:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.5)
        paragraph.paragraph_format.first_line_indent = Inches(-0.5)
        paragraph.add_run(ref)

    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = footer.add_run("Team MAIL – Hausarbeit Entwurf v1")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string("666666")

    OUT.parent.mkdir(exist_ok=True)
    doc.save(OUT)


if __name__ == "__main__":
    main()
