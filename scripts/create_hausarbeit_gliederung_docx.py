from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("docs/hausarbeit_gliederung.docx")


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(10)


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), "D0D7DE")
        borders.append(tag)
    tbl_pr.append(borders)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.add_run(item)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.add_run(item)


def add_callout(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_borders(table)
    cell = table.cell(0, 0)
    shade_cell(cell, "F4F6F9")
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(title)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string("1F3A5F")
    body_paragraph = cell.add_paragraph(body)
    body_paragraph.paragraph_format.space_after = Pt(0)
    doc.add_paragraph()


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for sec in doc.sections:
        sec.top_margin = Inches(1)
        sec.bottom_margin = Inches(1)
        sec.left_margin = Inches(1)
        sec.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_title(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run("Gliederung der Hausarbeit")
    run.font.name = "Calibri"
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("0B2545")

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    run = subtitle.add_run(
        "Team MAIL – Hybrid AI Agent for Dynamic Used Car Pricing | "
        "LNCS-Paper, 12 Seiten, Stand: 28.06.2026"
    )
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string("555555")


def main() -> None:
    doc = Document()
    configure_styles(doc)
    add_title(doc)

    add_callout(
        doc,
        "Ziel der Hausarbeit",
        "Die Arbeit soll das dreistufige Preissystem wissenschaftlich erklären: "
        "Stage 1 schätzt den Fahrzeugwert, Stage 2 passt das Marktpreisniveau "
        "über CPI an, Stage 3 ergänzt eine vorsichtige saisonale Korrektur. "
        "Die große Modellverbesserung gehört ausdrücklich zu Stage 1 V2, nicht zu Stage 3.",
    )

    doc.add_heading("Rahmenbedingungen aus dem Projektplan", level=1)
    add_bullets(
        doc,
        [
            "Abgabe: 31.07.2026.",
            "Umfang: ca. 12 Seiten im Springer-LNCS-Format.",
            "Geforderte Inhalte: Architektur, Methodik, Feature Engineering, Evaluation mit MAE/RMSE und wirtschaftliche bzw. strategische Interpretation.",
            "Endprodukt: reproduzierbare Python-Pipeline, Streamlit-Demo und wissenschaftliche Dokumentation.",
        ],
    )

    doc.add_heading("Empfohlene Hauptgliederung", level=1)
    sections = [
        (
            "1. Introduction",
            "Problemstellung, Motivation und Beitrag der Arbeit.",
            [
                "Warum Gebrauchtwagenpreise nicht statisch sind: Fahrzeugmerkmale, Marktpreisniveau, Saison.",
                "Ziel: dynamische, datenbasierte Preisempfehlung statt reinem Bauchgefühl.",
                "Beitrag: dreistufige Architektur mit ML-Basiswert, CPI-Marktupdate und Saisonregel.",
                "Kurz nennen: Streamlit-Demo als Proof of Concept.",
            ],
            "ca. 1 Seite",
        ),
        (
            "2. Data Sources and Preprocessing",
            "Beschreibung der Datenbasis und Aufbereitung.",
            [
                "Manheim/Kaggle Used Car Auction Prices: 558.743 US-Auktionsverkäufe 2014–2015.",
                "FRED-Makrodaten, insbesondere Used Cars and Trucks CPI (CUSR0000SETA01).",
                "Cleaning: fehlende Werte, Preis-/Kilometergrenzen, Datumsextraktion, Fahrzeugalter.",
                "Ausschlüsse: MMR, VIN und Seller nicht als Modellfeatures nutzen, um Leakage zu vermeiden.",
            ],
            "ca. 1,5 Seiten",
        ),
        (
            "3. System Architecture",
            "Gesamtarchitektur des Hybrid Agent.",
            [
                "Formel: Final Price = Stage-1-Basiswert × CPI-Multiplikator × Saisonfaktor.",
                "Stage Ownership sauber trennen: Stage 1 = Fahrzeugwert, Stage 2 = Markt, Stage 3 = Saison.",
                "Architekturdiagramm einbauen: Eingabe → Stage 1 → Stage 2 → Stage 3 → App/LLM-Erklärung.",
                "Begründen, warum ein hybrider Ansatz statt eines einzigen Black-Box-Modells gewählt wurde.",
            ],
            "ca. 1 Seite",
        ),
        (
            "4. Stage 1: Vehicle Value Model",
            "Kernmodell zur Schätzung des zeitneutralen Basispreises.",
            [
                "V1 kurz erklären: HistGradientBoosting als ursprüngliches Basismodell.",
                "V2 erklären: 50/50-Ensemble aus XGBoost raw target und log target.",
                "Features: model_year, vehicle_age, odometer, condition, make, model, trim, body, transmission, state, color, interior, make_model.",
                "Wichtig: V2 enthält bewusst keinen sale_month und kein year_month, damit Stage 2 und Stage 3 getrennt bleiben.",
                "Evaluation: V1 MAE 1.830,95 $, V2 MAE 1.370,15 $, Verbesserung 25,17 % auf gemeinsamem Split.",
            ],
            "ca. 2 Seiten",
        ),
        (
            "5. Stage 2: Macro-Level CPI Adjustment",
            "Marktpreisanpassung auf ein Zieljahr/einen Zielmonat.",
            [
                "CPI-Multiplikator relativ zum Jahresdurchschnitt 2015.",
                "Formel: Stage-2-Preis = Stage-1-Basiswert × CPI-Multiplikator.",
                "Aktueller Kontext: 2026-06 Multiplikator 1,2177, also ca. +21,8 % gegenüber 2015.",
                "Backtest einordnen: 1.370,16 $ → 1.376,22 $ MAE, nur +0,44 %, weil Testdaten nahe 2015 liegen.",
                "Nutzen liegt vor allem in Forward Projection für 2020–2026.",
            ],
            "ca. 1,25 Seiten",
        ),
        (
            "6. Stage 3: Seasonal Adjustment",
            "Saisonale Feinkorrektur nach Karosserieform und Monat.",
            [
                "Stage 3 baut kein neues Modell, sondern eine Regel auf bereinigten Residuen.",
                "Faktoren nach body und sale_month.",
                "Nicht rohe Monatsdurchschnitte verwenden, sondern CPI-normalisierte und fahrzeugmixbereinigte Abweichungen.",
                "Smoothing/Shrinkage Richtung 1,0; Faktorgrenzen 0,85 bis 1,15.",
                "Fehlende Monate August bis November bleiben neutral; Empfehlungen nur bei ausreichender Datenbasis.",
                "Evaluation: MAE 1.353,15 $ → 1.339,84 $, Verbesserung ca. 0,98 %.",
            ],
            "ca. 1,5 Seiten",
        ),
        (
            "7. Streamlit Prototype and User Interaction",
            "Beschreibung des lauffähigen Demonstrators.",
            [
                "Nutzer gibt Fahrzeugdaten ein; App berechnet finalen Verkaufspreis.",
                "Deutsche UI: Kilometer statt Meilen, deutsche Farben/Bundesstaaten/Begriffe.",
                "Validierung: Karosserieform passend zu Marke und Modell.",
                "Sterne-Skala für Zustand.",
                "Entwicklerdetails hinter Toggles, damit normale Nutzeroberfläche übersichtlich bleibt.",
            ],
            "ca. 1 Seite",
        ),
        (
            "8. Evaluation and Results",
            "Zusammenfassung der wichtigsten quantitativen Ergebnisse.",
            [
                "Stage 1: V2 verbessert MAE um 25,17 % gegenüber V1.",
                "Stage 2: historisch nahe 2015 nur geringe MAE-Veränderung, aber plausibler Marktpreisfaktor für spätere Jahre.",
                "Stage 3: kleine, aber methodisch saubere saisonale Verbesserung.",
                "Ergebnisse tabellarisch darstellen: MAE, RMSE, R², MAPE, Segmentergebnisse.",
                "Bootstrap-Intervall für Stage 1 erwähnen: 450,74 $ bis 470,83 $ MAE-Verbesserung.",
            ],
            "ca. 1,25 Seiten",
        ),
        (
            "9. Discussion and Limitations",
            "Kritische wissenschaftliche Einordnung.",
            [
                "US-Auktionsdaten statt Retailpreise; Ergebnisse nicht direkt auf EU-Markt übertragbar.",
                "Datensatz hauptsächlich 2014–2015; Saisonmonate August bis November fehlen.",
                "Random Split statt echter zeitlicher Zukunftstest.",
                "Stage 2 nutzt CPI als Gesamtmarktindikator, nicht modell- oder segmentgenau.",
                "Stage 3 ist bewusst konservativ und kein kausaler Nachweis.",
            ],
            "ca. 1 Seite",
        ),
        (
            "10. Conclusion and Outlook",
            "Abschluss und nächste Schritte.",
            [
                "Drei-Stufen-System liefert nachvollziehbare dynamische Preislogik.",
                "Stage 1 V2 bringt den größten Genauigkeitsgewinn.",
                "Stage 2 und 3 erhöhen Aktualität und Interpretierbarkeit.",
                "Ausblick: LLM-Orchestrierung, echter temporaler Holdout, weitere Märkte, Live-Daten, bessere Saisonabdeckung.",
            ],
            "ca. 0,75 Seiten",
        ),
    ]

    for title, purpose, bullets, pages in sections:
        doc.add_heading(title, level=2)
        paragraph = doc.add_paragraph()
        paragraph.add_run("Ziel des Kapitels: ").bold = True
        paragraph.add_run(purpose)
        paragraph = doc.add_paragraph()
        paragraph.add_run("Geplanter Umfang: ").bold = True
        paragraph.add_run(pages)
        add_bullets(doc, bullets)

    doc.add_heading("Welche Repo-Dateien für welches Kapitel nutzen?", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_borders(table)
    for cell, text in zip(table.rows[0].cells, ["Kapitel", "Primäre Repo-Quellen", "Wofür verwenden?"]):
        shade_cell(cell, "F2F4F7")
        set_cell_text(cell, text, bold=True)

    rows = [
        ("Data", "docs/data_cleaning.md; docs/feature_engineering.md; car_prices_clean.csv", "Datenbasis, Cleaning, Feature Engineering"),
        ("Stage 1", "docs/stage1/; scripts/train_stage1_v2.py; scripts/compare_stage1_v1_v2_shared_split.py", "Modellarchitektur, Features, V1/V2-Vergleich"),
        ("Stage 2", "docs/stage2/model_results_stage2.md; scripts/stage2_macro.py; macro_index.csv", "CPI-Logik, Marktmultiplikator, Backtest"),
        ("Stage 3", "docs/stage3/model_results_stage3.md; scripts/stage3_seasonality.py", "Saisonfaktoren, Datenabdeckung, Regel-Holdout"),
        ("App", "app/streamlit_app.py; scripts/stage1_runtime.py", "Integration, UX, Eingabevalidierung"),
        ("Projektlogik", "docs/stage_ownership.md; README.md; Aktueller_Stand.md", "Saubere Abgrenzung der Stages und Gesamtargumentation"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, row):
            set_cell_text(cell, text)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    doc.add_heading("Empfohlene Abbildungen und Tabellen", level=1)
    add_bullets(
        doc,
        [
            "Abbildung 1: Drei-Stufen-Architektur mit Formel Final Price = Stage 1 × Stage 2 × Stage 3.",
            "Tabelle 1: Übersicht der Features für Stage 1 V2.",
            "Tabelle 2: V1 vs. V2 Metriken auf identischem Split.",
            "Tabelle 3: CPI-Multiplikatoren für ausgewählte Jahre/Monate.",
            "Tabelle 4: Stage-3-Saisonfaktoren für ausgewählte Karosserieformen.",
            "Screenshot der Streamlit-App als Demonstrator, falls Platz vorhanden.",
        ],
    )

    doc.add_heading("Schreibpriorität für das Team", level=1)
    add_numbered(
        doc,
        [
            "Zuerst Methodik und Architektur schreiben, weil dort die Stage-Trennung sauber erklärt werden muss.",
            "Danach Evaluationstabellen einfügen und mit den Repo-Ergebnissen abgleichen.",
            "Erst am Ende Introduction und Conclusion final formulieren, damit sie zu den tatsächlichen Ergebnissen passen.",
            "Vor Abgabe prüfen, dass die 25,17 % klar als Stage-1-Ergebnis bezeichnet werden.",
            "Stage 3 nur als Saisonalität darstellen, nicht als neues Modell.",
        ],
    )
    add_callout(
        doc,
        "Kritischer Formulierungshinweis",
        "Nicht schreiben: „In Stage 3 wurde ein neues Modell gebaut.“ Besser: "
        "„Im Zuge der Integration wurde Stage 1 zu V2 verbessert; Stage 3 selbst bleibt die saisonale Korrekturschicht.“",
    )

    for sec in doc.sections:
        footer = sec.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = footer.add_run("Team MAIL – Hausarbeit Gliederung")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string("666666")

    OUT.parent.mkdir(exist_ok=True)
    doc.save(OUT)


if __name__ == "__main__":
    main()
