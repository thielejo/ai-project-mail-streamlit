from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("docs/hausarbeit_literatur_intro_theorie.docx")


def configure(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_title(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run("Erster Literaturentwurf: Einleitung und Theorieteil")
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("0B2545")

    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(14)
    run = sub.add_run(
        "Team MAIL – Hybrid AI Agent for Dynamic Used Car Pricing | Stand: 28.06.2026"
    )
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string("555555")


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(10)


def borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), "D0D7DE")
        tbl_borders.append(tag)
    tbl_pr.append(tbl_borders)


def add_callout(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    borders(table)
    cell = table.cell(0, 0)
    shade_cell(cell, "FFF7E6")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string("7A5A00")
    p2 = cell.add_paragraph(body)
    p2.paragraph_format.space_after = Pt(0)
    doc.add_paragraph()


def para(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run(text)


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_literature_table(doc: Document) -> None:
    doc.add_heading("Kurzüberblick der recherchierten Literatur", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    borders(table)
    headers = ["Quelle", "Kernaussage", "Nutzung in der Hausarbeit"]
    for cell, header in zip(table.rows[0].cells, headers):
        shade_cell(cell, "F2F4F7")
        set_cell(cell, header, bold=True)

    rows = [
        (
            "Akerlof (1970)",
            "Used-car markets sind ein klassisches Beispiel für Informationsasymmetrie und Qualitätsunsicherheit.",
            "Motivation: Warum datenbasierte Bewertung Vertrauen und Transparenz erhöhen kann.",
        ),
        (
            "Rosen (1974)",
            "Hedonic Pricing erklärt Preise als Bündel bewerteter Produkteigenschaften.",
            "Theoretische Grundlage für Fahrzeugfeatures wie Alter, Laufleistung, Zustand und Marke.",
        ),
        (
            "Chen & Guestrin (2016)",
            "XGBoost ist ein skalierbares Gradient-Boosting-System für tabellarische Daten.",
            "Begründung für Stage 1 V2 als starkes tabellarisches Regressionsmodell.",
        ),
        (
            "Pal et al. (2017)",
            "Random-Forest-basierte Used-Car-Preisprognose zeigt den Nutzen überwachter Lernverfahren.",
            "Einordnung unseres ML-Ansatzes in bestehende Used-Car-Pricing-Literatur.",
        ),
        (
            "Madhusudhanan et al. (2024)",
            "Neuere Used-Car-Pricing-Forschung betont Unsicherheit und vertrauenswürdige Prognosen.",
            "Argument für transparente Fehlerkennzahlen und vorsichtige Interpretierbarkeit.",
        ),
        (
            "BLS / FRED",
            "Used Cars and Trucks CPI misst Preisentwicklung gebrauchter Fahrzeuge als offiziellen Index.",
            "Theoretische und datenbezogene Begründung für Stage 2.",
        ),
    ]
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            set_cell(cell, value)


def main() -> None:
    doc = Document()
    configure(doc)
    add_title(doc)

    add_literature_table(doc)

    doc.add_heading("1. Erster Einleitungsversuch", level=1)
    para(
        doc,
        "Der Markt für Gebrauchtwagen ist durch eine Kombination aus hoher wirtschaftlicher Relevanz, "
        "starker Produktheterogenität und ausgeprägter Informationsasymmetrie geprägt. Bereits Akerlof "
        "(1970) nutzte den Gebrauchtwagenmarkt als zentrales Beispiel, um zu zeigen, wie Qualitätsunsicherheit "
        "zwischen Verkäufern und Käufern zu ineffizienten Marktmechanismen führen kann. Für Käufer ist der "
        "tatsächliche Zustand eines Fahrzeugs vor dem Kauf häufig nur unvollständig beobachtbar; für Verkäufer "
        "besteht gleichzeitig ein Anreiz, Fahrzeuge möglichst wertstabil darzustellen. Vor diesem Hintergrund "
        "gewinnen datenbasierte Bewertungsmodelle an Bedeutung, weil sie beobachtbare Fahrzeugmerkmale systematisch "
        "in eine transparente Preisindikation übersetzen können.",
    )
    para(
        doc,
        "Klassische Bewertungslogiken folgen dabei häufig einer hedonischen Perspektive: Der Preis eines Gutes "
        "ergibt sich nicht nur aus dem Gut als Ganzem, sondern aus einem Bündel einzelner Eigenschaften. Rosen "
        "(1974) beschreibt diesen Zusammenhang als Markt für implizite Preise, bei dem Produkteigenschaften wie "
        "Qualität, Ausstattung oder technische Merkmale in den beobachtbaren Marktpreis einfließen. Für Fahrzeuge "
        "bedeutet dies, dass Merkmale wie Marke, Modell, Fahrzeugalter, Laufleistung, Zustand, Karosserieform oder "
        "Ausstattung nicht isoliert betrachtet werden sollten, sondern gemeinsam den erwartbaren Wert bestimmen.",
    )
    para(
        doc,
        "Gleichzeitig reichen rein fahrzeugbezogene Merkmale nicht aus, um dynamische Gebrauchtwagenpreise vollständig "
        "zu erklären. Der US-Gebrauchtwagenmarkt unterliegt makroökonomischen Preisbewegungen, die sich unter anderem "
        "im Consumer Price Index für Used Cars and Trucks abbilden lassen. Die U.S. Bureau of Labor Statistics (n.d.) "
        "beschreibt diesen Index als Teil des Consumer Price Index, der gebrauchte Fahrzeuge im Alter von zwei bis "
        "sieben Jahren umfasst. Über FRED wird dieser Index als monatliche Zeitreihe bereitgestellt (U.S. Bureau of "
        "Labor Statistics, 2026). Ein praxistaugliches Preissystem muss daher sowohl den fahrzeugspezifischen Basiswert "
        "als auch das aktuelle Marktpreisniveau berücksichtigen.",
    )
    para(
        doc,
        "Die vorliegende Arbeit entwickelt deshalb einen hybriden AI Agent zur dynamischen Gebrauchtwagenpreisgestaltung. "
        "Das System kombiniert ein maschinelles Lernmodell für den fahrzeugbezogenen Basiswert mit einer makroökonomischen "
        "CPI-Anpassung und einer saisonalen Feinkorrektur. Stage 1 schätzt den zeitneutralen Fahrzeugwert auf Basis "
        "tabellarischer Fahrzeugmerkmale. Stage 2 aktualisiert diesen Wert über einen CPI-Multiplikator auf das gewünschte "
        "Marktpreisniveau. Stage 3 ergänzt schließlich eine vorsichtige Saisonalitätsregel nach Karosserieform und Monat. "
        "Damit verbindet der Ansatz die Prognosekraft moderner Machine-Learning-Verfahren mit der Interpretierbarkeit "
        "regelbasierter ökonomischer Korrekturen.",
    )
    para(
        doc,
        "Der Beitrag der Arbeit liegt in einer reproduzierbaren Drei-Stufen-Architektur für dynamische Fahrzeugbewertung. "
        "Insbesondere zeigt das Projekt, dass ein verbessertes Stage-1-Modell auf Basis eines XGBoost-Ensembles den "
        "mittleren absoluten Fehler gegenüber dem ursprünglichen Modell deutlich reduzieren kann. Dieser Genauigkeitsgewinn "
        "ist jedoch klar von Stage 3 zu trennen: Die große Modellverbesserung gehört zum Fahrzeugwertmodell in Stage 1, "
        "während Stage 3 ausschließlich eine kleine, konservative Saisonkorrektur darstellt. Diese Trennung ist zentral, "
        "um die Ergebnisse methodisch korrekt zu interpretieren.",
    )

    doc.add_heading("2. Erster Theorieteil", level=1)
    doc.add_heading("2.1 Informationsasymmetrie im Gebrauchtwagenmarkt", level=2)
    para(
        doc,
        "Der Gebrauchtwagenmarkt ist ein klassisches Beispiel für asymmetrische Information. Verkäufer kennen die Historie, "
        "Pflege und mögliche verborgene Mängel eines Fahrzeugs häufig besser als potenzielle Käufer. Akerlof (1970) "
        "zeigt am Beispiel sogenannter „lemons“, dass solche Informationsvorteile zu adverser Selektion führen können: "
        "Wenn Käufer die tatsächliche Qualität nicht zuverlässig erkennen, orientieren sie sich am erwarteten Durchschnitt. "
        "Dadurch kann der Marktpreis für hochwertige Fahrzeuge zu niedrig werden, während minderwertige Fahrzeuge relativ "
        "attraktiv zu verkaufen sind. Datenbasierte Preismodelle können dieses Problem nicht vollständig lösen, aber sie "
        "reduzieren Unsicherheit, indem sie beobachtbare Eigenschaften konsistent und nachvollziehbar bewerten.",
    )

    doc.add_heading("2.2 Hedonic Pricing als Grundlage der Fahrzeugbewertung", level=2)
    para(
        doc,
        "Die hedonische Preistheorie bildet eine zentrale theoretische Grundlage für tabellarische Fahrzeugbewertung. "
        "Nach Rosen (1974) spiegeln Marktpreise die impliziten Werte einzelner Produktattribute wider. Ein Fahrzeugpreis "
        "kann daher als Ergebnis mehrerer Merkmalsbeiträge verstanden werden, etwa Fahrzeugalter, Laufleistung, Zustand, "
        "Marke, Modell oder Ausstattung. Dieser Gedanke passt unmittelbar zu Stage 1 des Projekts: Das Modell versucht, "
        "aus beobachtbaren Fahrzeugeigenschaften einen Basiswert zu schätzen. Moderne Machine-Learning-Verfahren erweitern "
        "diese Logik, indem sie nichtlineare Zusammenhänge und Interaktionen zwischen Merkmalen abbilden können.",
    )

    doc.add_heading("2.3 Machine Learning für Used-Car-Pricing", level=2)
    para(
        doc,
        "Used-Car-Pricing ist ein typisches Regressionsproblem auf tabellarischen Daten. Frühere Arbeiten zeigen, dass "
        "überwachte Lernverfahren wie Random Forests für die Preisprognose gebrauchter Fahrzeuge eingesetzt werden können "
        "(Pal et al., 2017). Neuere Forschung betrachtet darüber hinaus probabilistische Ansätze, um nicht nur Punktpreise, "
        "sondern auch Prognoseunsicherheit abzubilden (Madhusudhanan et al., 2024). Für das vorliegende Projekt ist "
        "insbesondere Gradient Boosting relevant. Chen und Guestrin (2016) beschreiben XGBoost als skalierbares Tree-Boosting-"
        "System, das sich für große tabellarische Datensätze eignet und in vielen Machine-Learning-Anwendungen starke "
        "Ergebnisse erzielt. Das Stage-1-V2-Modell knüpft an diese Literatur an, indem es ein XGBoost-Ensemble zur "
        "Schätzung des fahrzeugbezogenen Basispreises verwendet.",
    )

    doc.add_heading("2.4 Makroökonomische Preisindizes und CPI-Anpassung", level=2)
    para(
        doc,
        "Während Stage 1 den relativen Fahrzeugwert aus Merkmalen ableitet, adressiert Stage 2 die Veränderung des allgemeinen "
        "Marktpreisniveaus. Der Consumer Price Index für Used Cars and Trucks ist hierfür ein geeigneter externer Indikator, "
        "weil er die Preisentwicklung gebrauchter Fahrzeuge in den USA als offizielle Zeitreihe abbildet. Die BLS-Factsheet-"
        "Beschreibung grenzt den Index auf gebrauchte Fahrzeuge im Alter von zwei bis sieben Jahren ein und erläutert seine "
        "Einbettung in den Consumer Price Index (U.S. Bureau of Labor Statistics, n.d.). Über FRED steht die Reihe als "
        "monatliche, saisonbereinigte Indexreihe zur Verfügung (U.S. Bureau of Labor Statistics, 2026). Methodisch wird "
        "Stage 2 daher als Multiplikator modelliert: Ein zeitneutraler Stage-1-Basiswert wird auf das Preisniveau eines "
        "gewählten Zielmonats skaliert.",
    )

    doc.add_heading("2.5 Saisonalität und regelbasierte Korrekturen", level=2)
    para(
        doc,
        "Neben langfristigen Marktpreisbewegungen können auch saisonale Nachfrageeffekte eine Rolle spielen. Für Fahrzeuge "
        "ist plausibel, dass bestimmte Karosserieformen in einzelnen Monaten stärker oder schwächer nachgefragt werden, "
        "beispielsweise Cabriolets in wärmeren Monaten oder SUVs in Phasen erhöhter Wintervorbereitung. Aus methodischer "
        "Sicht ist jedoch entscheidend, solche Effekte nicht aus rohen Monatsdurchschnitten abzuleiten, weil sich der "
        "Fahrzeugmix zwischen Monaten unterscheiden kann. Stage 3 nutzt deshalb keine neue Black-Box-Modellierung, sondern "
        "berechnet konservative Saisonfaktoren aus CPI-normalisierten und fahrzeugmixbereinigten Modellabweichungen. Dadurch "
        "bleibt die Saisonalitätslogik interpretierbar und wird klar von der eigentlichen Fahrzeugwertschätzung getrennt.",
    )

    doc.add_heading("2.6 Hybride AI-Agent-Architektur", level=2)
    para(
        doc,
        "Die Kombination aus Machine Learning, ökonomischem Index und regelbasierter Saisonalität entspricht einer hybriden "
        "Agentenarchitektur. Ein rein datengetriebenes End-to-End-Modell könnte zwar alle verfügbaren Informationen gemeinsam "
        "verarbeiten, wäre aber schwerer zu interpretieren und könnte Markt-, Fahrzeug- und Saisonanteile vermischen. Die "
        "dreistufige Architektur trennt diese Ebenen bewusst: Stage 1 erklärt den Fahrzeugwert, Stage 2 das Marktpreisniveau "
        "und Stage 3 die saisonale Feinkorrektur. Für eine wissenschaftliche Arbeit ist diese Trennung besonders hilfreich, "
        "weil jede Komponente separat begründet, getestet und kritisch diskutiert werden kann.",
    )

    add_callout(
        doc,
        "Methodischer Prüfhinweis für das Team",
        "Bei der Literatur- und Datenprüfung fiel auf: FRED führt „Used Cars and Trucks“ unter CUSR0000SETA02. "
        "CUSR0000SETA01 ist auf FRED „New Vehicles“. Bitte vor finaler Paper-Abgabe prüfen, ob Code, Dokumentation "
        "und Paper denselben korrekten CPI-Series-Code verwenden.",
    )

    doc.add_heading("Literaturverzeichnis nach APA 7", level=1)
    refs = [
        "Akerlof, G. A. (1970). The market for “lemons”: Quality uncertainty and the market mechanism. The Quarterly Journal of Economics, 84(3), 488–500. https://doi.org/10.2307/1879431",
        "Chen, T., & Guestrin, C. (2016). XGBoost: A scalable tree boosting system. In Proceedings of the 22nd ACM SIGKDD International Conference on Knowledge Discovery and Data Mining (pp. 785–794). Association for Computing Machinery. https://doi.org/10.1145/2939672.2939785",
        "Madhusudhanan, K., Behrens, G., Stubbemann, M., & Schmidt-Thieme, L. (2024). ProbSAINT: Probabilistic tabular regression for used car pricing. arXiv. https://arxiv.org/abs/2403.03812",
        "Pal, N., Arora, P., Sundararaman, D., Kohli, P., & Palakurthy, S. S. (2017). How much is my car worth? A methodology for predicting used cars prices using Random Forest. arXiv. https://arxiv.org/abs/1711.06970",
        "Rosen, S. (1974). Hedonic prices and implicit markets: Product differentiation in pure competition. Journal of Political Economy, 82(1), 34–55. https://doi.org/10.1086/260169",
        "U.S. Bureau of Labor Statistics. (n.d.). Measuring price change in the CPI: Used cars and trucks. U.S. Department of Labor. https://www.bls.gov/cpi/factsheets/used-cars-and-trucks.htm",
        "U.S. Bureau of Labor Statistics. (2026). Consumer Price Index for All Urban Consumers: Used Cars and Trucks in U.S. City Average [CUSR0000SETA02]. FRED, Federal Reserve Bank of St. Louis. https://fred.stlouisfed.org/series/CUSR0000SETA02",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        p.add_run(ref)

    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = footer.add_run("Team MAIL – Literaturentwurf")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string("666666")

    OUT.parent.mkdir(exist_ok=True)
    doc.save(OUT)


if __name__ == "__main__":
    main()
